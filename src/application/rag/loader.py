import re
from pathlib import Path
from langchain_core.documents import Document
from langchain_community.document_loaders import TextLoader
import fitz  # PyMuPDF

# Docling Imports
from docling.document_converter import DocumentConverter, PdfFormatOption
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import PdfPipelineOptions

# -------------------------
# 1. GLOBAL CONVERTER (SPEED FIX)
# -------------------------
_DOCLING_CONVERTER = None

def get_docling_converter(enable_ocr: bool = True):
    """Singleton pattern agar model Docling diload hanya 1 kali di memori."""
    global _DOCLING_CONVERTER
    if _DOCLING_CONVERTER is None:
        print("⚙️ Inisialisasi Model Docling (Hanya 1 kali, mohon tunggu)...")
        pipeline_options = PdfPipelineOptions()
        pipeline_options.do_ocr = enable_ocr  
        pipeline_options.do_table_structure = True 
        pipeline_options.table_structure_options.do_cell_matching = True

        _DOCLING_CONVERTER = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
            }
        )
    return _DOCLING_CONVERTER

# -------------------------
# 2. EXTRACTORS & UTILITIES
# -------------------------
def is_gibberish(text: str) -> bool:
    """
    Mendeteksi apakah teks rusak (corrupt font / bad OCR / Mojibake).
    Ciri-cirinya: terlalu banyak karakter aneh atau kurangnya huruf standar.
    """
    if not text or len(text) < 50:
        return True
        
    # Hitung jumlah huruf dan angka standar
    alphanumeric_count = len(re.findall(r'[a-zA-Z0-9]', text))
    alphanumeric_ratio = alphanumeric_count / len(text)
    
    # Hitung jumlah simbol aneh (selain spasi, huruf, angka, dan tanda baca umum)
    weird_symbols_count = len(re.findall(r'[^a-zA-Z0-9\s\.,\-\?!\(\)]', text))
    weird_symbols_ratio = weird_symbols_count / len(text)
    
    # Kriteria Gibberish:
    # 1. Jika kurang dari 65% teks adalah huruf/angka, ATAU
    # 2. Jika lebih dari 5% teks adalah simbol aneh
    if alphanumeric_ratio < 0.65 or weird_symbols_ratio > 0.05:
        return True
        
    return False

def extract_with_pymupdf(pdf_path: str) -> str:
    """Ekstraksi super cepat untuk PDF berbasis teks (digital born)."""
    try:
        doc = fitz.open(pdf_path)
        texts = []
        for page in doc:
            text = page.get_text("text")
            if text:
                texts.append(text)
        return "\n".join(texts).strip()
    except Exception as e:
        print(f"⚠️ PyMuPDF Error on {pdf_path}: {e}")
        return ""

def extract_with_docling(pdf_path: str) -> str:
    """Ekstraksi cerdas. Menggunakan Global Converter."""
    try:
        converter = get_docling_converter()
        result = converter.convert(pdf_path)
        
        if not result or not result.document:
            return ""

        return result.document.export_to_markdown().strip()
    except Exception as e:
        print(f"❌ Docling Error on {pdf_path}: {e}")
        return ""

def extract_word(docx_path: str) -> str:
    """Ekstraksi teks dari file Word (.docx) menggunakan python-docx."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(docx_path)
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts).strip()
    except ImportError:
        print("❌ python-docx not installed. Run: pip install python-docx")
        return ""
    except Exception as e:
        print(f"❌ Word extraction error on {docx_path}: {e}")
        return ""

def extract_image(image_path: str) -> str:
    """Ekstraksi teks dari file gambar menggunakan pytesseract OCR/Docling."""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img, lang="eng+ind").strip()
        if text:
            print(f"   ✅ OCR via pytesseract ({len(text)} chars)")
            return text
        print("   ⚠️ pytesseract returned empty text — trying Docling fallback...")
    except ImportError:
        print("   ⚠️ pytesseract/Pillow not installed — trying Docling fallback...")
    except Exception as e:
        print(f"   ⚠️ pytesseract error: {e} — trying Docling fallback...")

    try:
        converter = get_docling_converter()
        result = converter.convert(image_path)
        if result and result.document:
            text = result.document.export_to_markdown().strip()
            if text:
                print(f"   ✅ OCR via Docling ({len(text)} chars)")
                return text
    except Exception as e:
        print(f"   ❌ Docling image error: {e}")

    return ""

# -------------------------
# SUPPORTED EXTENSIONS
# -------------------------
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"}
WORD_EXTENSIONS  = {".docx", ".doc"}

# -------------------------
# 3. MAIN LOADER — load a single file
# -------------------------
def load_single_document(file_path: str, min_text_length: int = 300):
    """
    Load and extract text from a single file.
    Returns a LangChain Document.
    """
    file = Path(file_path)
    suffix = file.suffix.lower()

    # -------- PDF --------
    if suffix == ".pdf":
        print(f"📄 Processing PDF: {file.name}...")
        
        # 1. Selalu jalankan PyMuPDF pertama untuk kecepatan maksimal
        final_text = extract_with_pymupdf(str(file))
        loader_used = "pymupdf"

        # 2. Evaluasi: Apakah teks terlalu sedikit ATAU terdeteksi sebagai gibberish?
        if len(final_text) < min_text_length or is_gibberish(final_text):
            if is_gibberish(final_text) and len(final_text) > 0:
                print(f"   ⚠️ Terdeteksi teks Gibberish/Corrupt ({len(final_text)} chars). Memaksa Docling OCR...")
            else:
                print(f"   ⚠️ Teks terlalu sedikit ({len(final_text)} chars). Terdeteksi PDF Scan. Pindah ke Docling OCR...")
                
            # Fallback ke Docling OCR
            ocr_text = extract_with_docling(str(file))
            
            # Gunakan hasil Docling asalkan dia mengembalikan sesuatu yang cukup
            if len(ocr_text.strip()) > 50:
                final_text = ocr_text
                loader_used = "docling_ocr"

        if not final_text.strip():
            print(f"   ❌ Failed to extract: {file.name}")
            return None

        print(f"   ✅ Done ({loader_used}). Length: {len(final_text)} chars.")
        
        metadata = {"source": file.name, "file_type": "pdf", "loader": loader_used}
        return Document(page_content=final_text, metadata=metadata)

    # -------- TXT --------
    elif suffix == ".txt":
        try:
            docs = TextLoader(str(file)).load()
            if docs:
                docs[0].metadata.update({"file_type": "txt"})
                return docs[0]
        except Exception as e:
            print(f"❌ Failed to load TXT {file.name}: {e}")
        return None

    # -------- WORD (.docx / .doc) --------
    elif suffix in WORD_EXTENSIONS:
        print(f"📝 Processing Word file: {file.name}...")
        text = extract_word(str(file))
        if not text.strip():
            print(f"   ❌ Failed to extract: {file.name}")
            return None
            
        print(f"   ✅ Done (python-docx). Length: {len(text)} chars.")
        
        metadata = {"source": file.name, "file_type": "docx", "loader": "python-docx"}
        return Document(page_content=text, metadata=metadata)

    # -------- IMAGE --------
    elif suffix in IMAGE_EXTENSIONS:
        print(f"🖼️ Processing image: {file.name}...")
        text = extract_image(str(file))
        if not text.strip():
            print(f"   ❌ No text extracted from image: {file.name}")
            return None
            
        print(f"   ✅ Done (OCR). Length: {len(text)} chars.")
        
        metadata = {"source": file.name, "file_type": "image", "loader": "ocr"}
        return Document(page_content=text, metadata=metadata)

    else:
        print(f"⏭️ Skipping unsupported file type: {file.name}")
        return None

# -------------------------
# MAIN LOADER — load all docs in a directory
# -------------------------
def load_documents(docs_path: str, min_text_length: int = 300):
    documents = []
    path = Path(docs_path)

    if not path.exists():
        print(f"⚠️ Path not found: {docs_path}")
        return documents

    for file in path.iterdir():
        if not file.is_file():
            continue
        doc = load_single_document(str(file), min_text_length)
        if doc:
            documents.append(doc)

    return documents